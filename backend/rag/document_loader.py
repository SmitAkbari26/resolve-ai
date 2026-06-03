"""
Resolve AI - Document Loader
Loads documents from PDFs, Markdown, DOCX, and plain text files.
"""

from pathlib import Path
from typing import Generator
import json

from utils.logger import get_logger
from utils.text_sanitize import sanitize_text

logger = get_logger(__name__)


class Document:
    """Represents a loaded document with content and metadata."""

    def __init__(self, content: str, metadata: dict | None = None):
        self.content = sanitize_text(content)
        self.metadata = metadata or {}

    def __repr__(self):
        source = self.metadata.get("source", "unknown")
        return f"Document(source={source}, length={len(self.content)})"


def load_pdf(file_path: Path) -> Document:
    """Load text content from a PDF file."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return Document(
            content=text.strip(),
            metadata={"source": str(file_path), "type": "pdf", "pages": len(reader.pages)},
        )
    except Exception as e:
        logger.error(f"Failed to load PDF {file_path}: {e}")
        return Document(content="", metadata={"source": str(file_path), "error": str(e)})


def load_markdown(file_path: Path) -> Document:
    """Load text content from a Markdown file."""
    text = file_path.read_text(encoding="utf-8")
    return Document(
        content=text.strip(),
        metadata={"source": str(file_path), "type": "markdown"},
    )


def load_docx(file_path: Path) -> Document:
    """Load text content from a DOCX file."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(file_path))
        text = "\n".join(para.text for para in doc.paragraphs)
        return Document(
            content=text.strip(),
            metadata={"source": str(file_path), "type": "docx"},
        )
    except Exception as e:
        logger.error(f"Failed to load DOCX {file_path}: {e}")
        return Document(content="", metadata={"source": str(file_path), "error": str(e)})


def load_text(file_path: Path) -> Document:
    """Load text content from a plain text file."""
    text = file_path.read_text(encoding="utf-8")
    return Document(
        content=text.strip(),
        metadata={"source": str(file_path), "type": "text"},
    )


def load_json(file_path: Path) -> Document:
    """Load text content from a JSON file."""
    try:
        data = json.loads(file_path.read_text(encoding="utf-8"))
        # Format JSON as a pretty-printed string for embedding
        text = json.dumps(data, indent=2)
        return Document(
            content=text.strip(),
            metadata={"source": str(file_path), "type": "json"},
        )
    except Exception as e:
        logger.error(f"Failed to load JSON {file_path}: {e}")
        return Document(content="", metadata={"source": str(file_path), "error": str(e)})


# File extension → loader mapping
LOADERS = {
    ".pdf": load_pdf,
    ".md": load_markdown,
    ".docx": load_docx,
    ".txt": load_text,
    ".json": load_json,
}


def load_file(file_path: Path) -> Document | None:
    """Load a single file by extension."""
    path = Path(file_path)
    if not path.is_file():
        logger.warning(f"File not found: {path}")
        return None

    suffix = path.suffix.lower()
    loader = LOADERS.get(suffix)
    if not loader:
        logger.warning(f"Unsupported file type: {suffix}")
        return None

    doc = loader(path)
    if not doc.content:
        logger.warning(f"Empty content after load: {path}")
        return None
    return doc


def load_documents(directory: Path) -> list[Document]:
    """
    Load all supported documents from a directory (recursively).
    Supports: .pdf, .md, .docx, .txt
    """
    documents = []

    if not directory.exists():
        logger.warning(f"Directory does not exist: {directory}")
        return documents

    for file_path in directory.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() in LOADERS:
            loader = LOADERS[file_path.suffix.lower()]
            doc = loader(file_path)
            if doc.content:
                documents.append(doc)
                logger.info(f"Loaded: {file_path.name} ({len(doc.content)} chars)")

    logger.info(f"Total documents loaded: {len(documents)}")
    return documents
